
import pdfplumber
import docx
import re
# from pyresparser import ResumeParser
from datetime import datetime
import os
from sqlalchemy.ext.asyncio import AsyncSession
from typing import List
from google.generativeai import GenerativeModel, configure
from dotenv import load_dotenv
from models import ResumeLog
from sqlalchemy.ext.asyncio import AsyncSession


# Load environment variables
load_dotenv()
configure(api_key=os.getenv("GEMINI_API_KEY"))

# Initialize Gemini model
gemini = GenerativeModel("gemini-2.5-pro")

# Role-specific required skills
ROLE_SKILLS = {
    "Machine Learning Engineer": {"Python", "NumPy", "Pandas", "Scikit-learn", "TensorFlow", "PyTorch"},
    "Frontend Developer": {"HTML", "CSS", "JavaScript", "React.js", "UI/UX"},
    "Backend Developer": {"Python", "SQL", "MongoDB", "FastAPI", "Django"},
}

# Extract text from resume
def extract_text(file_path):
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return ""

# Extract email from resume
def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return match.group(0) if match else None

# Fallback skill extraction
def extract_skills_fallback(text, known_skills: List[str]):
    resume_words = set(re.findall(r'\b\w[\w+\.\#]*\b', text.lower()))
    found = [skill for skill in known_skills if skill.lower() in resume_words]
    return list(set(found))

# Experience estimation
def adjust_experience(exp_years, text):
    try:
        exp_years = float(exp_years)
    except:
        exp_years = 0
    if exp_years == 0 and any(word in text.lower() for word in ["intern", "training", "club"]):
        exp_years = 0.5
    return exp_years

def get_experience_level(exp_years):
    if exp_years >= 3:
        return "senior"
    elif exp_years >= 1:
        return "mid"
    else:
        return "junior"

# Parse resume
# def parse_resume(file_path):
#     try:
#         return ResumeParser(file_path).get_extracted_data()
#     except:
#         return {}
import re
import docx
import pdfplumber

def parse_resume(file_path):
    # Extract raw text
    text = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs)

    # Try basic name extraction: first non-empty short line
    name = None
    for line in text.strip().splitlines():
        line = line.strip()
        if line and len(line.split()) <= 4:
            name = line
            break
    if not name:
        name = "Candidate"

    # Extract email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    email = email_match.group(0) if email_match else None

    return {
        "name": name,
        "email": email,
        "skills": [],  # You already have fallback skill extractor
        "total_experience": 0  # You already adjust and estimate this separately
    }


# Match skills to role
def compute_skill_match(resume_skills, required_skills):
    matched = set(resume_skills or []) & set(required_skills)
    return len(matched) / len(required_skills) if required_skills else 0, list(matched)

# Gemini LLM scoring
def get_gemini_score(resume_text, job_title, job_description):
    prompt = f"""
    Given the following resume:
    {resume_text}

    Evaluate how well this resume matches the job titled '{job_title}' with the following description:
    {job_description}

    Return only a numeric ATS score between 0 and 1, where 1 means a perfect match.
    """
    try:
        response = gemini.generate_content(prompt)
        score = float(re.findall(r"\d+\.\d+", response.text)[0])
        return min(max(score, 0), 1)
    except:
        return 0.5



async def get_required_skills_from_llm(job_title: str, job_description: str = "") -> List[str]:
    prompt = f"""
    Given the job title "{job_title}" and the following description:

    {job_description if job_description else '[No additional description provided]'}

    List 8 to 12 important skills (technical and soft) required for this role.
    Return the list as bullet points, one per line.
    """
    try:
        response = gemini.generate_content(prompt)
        print("⚡ Gemini response:", response.text)
        lines = response.text.splitlines()
        skills = [line.strip("-• ").strip() for line in lines if line.strip()]
        # print(skills)
        return skills
    except Exception as e:
        print(f"❌ Gemini skill generation failed: {e}")
        return []

# Main analysis function
async def analyze_resume(file_path, job_title, job_description, job_id, thresholds, db: AsyncSession,required_skills=None):
    resume_text = extract_text(file_path)
    data = parse_resume(file_path)

    name = data.get("name")
    if not name or len(name.strip()) < 2:
        first_lines = resume_text.strip().splitlines()[:5]
        for line in first_lines:
            if line.strip() and len(line.split()) <= 4:
                name = line.strip()
                break
        if not name:
            name = "Candidate"

    email = data.get("email") or extract_email(resume_text)
    skills = data.get("skills", [])
    # required_skills = ROLE_SKILLS.get(job_title, set())
    required_skills = required_skills.split(",") if required_skills else ROLE_SKILLS.get(job_title, set())
    if not skills:
        skills = extract_skills_fallback(resume_text, required_skills)

    raw_exp = data.get("total_experience", 0)
    exp_years = adjust_experience(raw_exp, resume_text)
    level = get_experience_level(exp_years)

    llm_score = get_gemini_score(resume_text, job_title, job_description)
    skill_score, matched_skills = compute_skill_match(skills, required_skills)

    final_score = 0.7 * llm_score + 0.3 * skill_score

    threshold_map = thresholds or {"junior": 0.45, "mid": 0.55, "senior": 0.6}
    threshold = threshold_map.get(level, 0.5)
    status = "ACCEPTED" if final_score >= threshold else "REJECTED"

    # Store in database
    log = ResumeLog(
        name=name,
        email=email,
        role=job_title,
        experience_level=level,
        final_score=round(final_score, 2),
        status=status,
        job_id=job_id
    )
    db.add(log)
    await db.commit()

    return {
        "name": name,
        "email": email,
        "job_title": job_title,
        "experience_years": exp_years,
        "level": level,
        "llm_score": round(llm_score, 2),
        "skill_score": round(skill_score, 2),
        "final_score": round(final_score, 2),
        "status": status,
        "matched_skills": matched_skills,
        "required_skills": list(required_skills),
        "job_id": job_id
    }
